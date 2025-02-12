import subprocess
from typing import List, Tuple

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml

def ensure_hyperlink_style_exists(doc: Document) -> None:
    """
    Create a dummy paragraph with a run that uses the 'Hyperlink' style,
    so that the style is forced to appear in the document.
    """
    try:
        para = doc.add_paragraph()
        run = para.add_run("dummy", style="Hyperlink")
    except KeyError:
        print("[INFO] 'Hyperlink' style not found by dummy run; will add it manually.")

def add_hyperlink_style(doc):
    """
    Manually add a Hyperlink character style to the document.
    """
    hyperlink_style_xml = r'''
    <w:style w:type="character" w:styleId="Hyperlink"
      xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:name w:val="Hyperlink"/>
      <w:basedOn w:val="DefaultParagraphFont"/>
      <w:link w:val="DefaultParagraphFont"/>
      <w:uiPriority w:val="99"/>
      <w:unhideWhenUsed/>
      <w:qFormat/>
      <w:rPr>
        <w:color w:val="0000FF"/>
        <w:u w:val="single"/>
      </w:rPr>
    </w:style>
    '''
    styles_element = doc.styles.element
    styles_element.append(parse_xml(hyperlink_style_xml))
 

def override_hyperlink_style(
    doc: Document,
    custom_name: str = "IBM Hyperlink",
    font_size: float = 11,
    underline: bool = True,
    font_color: Tuple[int, int, int] = (0, 0, 255)
) -> None:
    """
    Override the built-in 'Hyperlink' character style.
    """
    try:
        style = doc.styles["Hyperlink"]
    except KeyError:
        print("[WARNING] 'Hyperlink' style not found. Skipping override.")
        return

    if style.type != WD_STYLE_TYPE.CHARACTER:
        print("[WARNING] 'Hyperlink' style is not a character style. Skipping override.")
        return

    # Change the style's name if desired (this is optional)
    original_id = style.style_id
    style.name = custom_name
    style.style_id = original_id

    font = style.font
    font.size = Pt(font_size)
    font.underline = underline
    r, g, b = font_color
    font.color.rgb = RGBColor(r, g, b)


def ensure_built_in_styles_exist(doc: DocxDocument, style_names: List[str]) -> None:
    """
    Add a temporary paragraph for each built-in *paragraph* style so they appear (not latent).
    """
    for style_name in style_names:
        print(f"[INFO] Ensuring paragraph style '{style_name}' exists")
        doc.add_paragraph("X", style=style_name)


def remove_temp_paragraphs(doc: DocxDocument) -> None:
    """
    Remove the dummy paragraphs that contain only 'X'.
    """
    to_remove = []
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "X":
            to_remove.append(i)
    for idx in reversed(to_remove):
        p = doc.paragraphs[idx]
        parent = p._element.getparent()
        parent.remove(p._element)


def override_built_in_style(
    doc: DocxDocument,
    base_name: str,
    custom_name: str,
    font_size: float,
    bold: bool,
    italic: bool,
    font_color: Tuple[int, int, int],
    space_before: float,
    space_after: float
) -> None:
    """
    Override a built-in paragraph style in doc.styles[base_name], 
    but keep the original style ID (so Pandoc recognizes it).
    """
    try:
        style = doc.styles[base_name]
    except KeyError:
        print(f"[WARNING] style '{base_name}' not found. Skipping override.")
        return

    if style.type != WD_STYLE_TYPE.PARAGRAPH:
        print(f"[WARNING] style '{base_name}' is not a paragraph style. Skipping.")
        return

    original_id = style.style_id
    style.name = custom_name
    style.style_id = original_id

    # Font settings
    font = style.font
    font.name = "IBM Plex Sans"  # Or any font you want

    # Remove theme attributes
    rPr = font.element.rPr
    if rPr is not None and rPr.rFonts is not None:
        for theme_attrib in ("w:asciiTheme", "w:hAnsiTheme", "w:csTheme", "w:eastAsiaTheme"):
            rPr.rFonts.attrib.pop(qn(theme_attrib), None)

    font.size = Pt(font_size)
    font.bold = bold
    font.italic = italic

    # Color
    r, g, b = font_color
    font.color.rgb = RGBColor(r, g, b)

    # Spacing
    pf = style.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)


def create_reference_doc(config: dict, reference_docx: str) -> None:
    """
    1) Create a new .docx
    2) Force built-in paragraph styles to appear
    3) Override each style from the YAML config
    4) Remove dummy paragraphs
    5) Save as reference_docx (for Pandoc's --reference-doc)
    """
    doc = Document()

    # Add paragraph styles so they appear
    needed_paragraph_styles = [
        "Normal", "Title", "Subtitle",
        "Heading 1", "Heading 2", "Heading 3",
        "Heading 4", "Heading 5", "Heading 6",
        "Quote"
    ]
    ensure_built_in_styles_exist(doc, needed_paragraph_styles)

    # Read style definitions from config
    for style_def in config.get("styles", []):
        base_name = style_def["base_name"]
        custom_name = style_def["custom_name"]
        font_size = style_def["font_size"]
        bold = style_def.get("bold", False)
        italic = style_def.get("italic", False)
        font_color = tuple(style_def.get("font_color", [0, 0, 0]))
        space_before = style_def.get("space_before", 0)
        space_after = style_def.get("space_after", 0)

        override_built_in_style(
            doc,
            base_name=base_name,
            custom_name=custom_name,
            font_size=font_size,
            bold=bold,
            italic=italic,
            font_color=font_color,
            space_before=space_before,
            space_after=space_after
        )

    # Remove dummy paragraphs
    ensure_hyperlink_style_exists(doc)
    add_hyperlink_style(doc)
    override_hyperlink_style(doc)
    remove_temp_paragraphs(doc)

    # Save
    doc.save(reference_docx)
    print(f"[INFO] Created reference DOCX: {reference_docx}")


def convert_md_to_word(
    input_md: str,
    output_docx: str,
    reference_docx: str,
    from_format: str = "markdown+footnotes+mark"
) -> None:
    """
    Use Pandoc to convert Markdown -> Word docx,
    applying the specified reference doc.
    """
    cmd = [
        "pandoc",
        f"--from={from_format}",
        input_md,
        f"--reference-doc={reference_docx}",
        "-o",
        output_docx
    ]
    print(f"[INFO] Running Pandoc: {' '.join(cmd)}")
    subprocess.run(cmd, check=True)
    print(f"[INFO] Converted {input_md} -> {output_docx}")
