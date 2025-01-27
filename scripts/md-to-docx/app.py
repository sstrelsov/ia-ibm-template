#!/usr/bin/env python3

import subprocess
import sys
from typing import List, Tuple

import yaml
from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.styles.style import BaseStyle
from docx.table import Table
from styles import apply_table_style


def ensure_built_in_styles_exist(doc: DocxDocument, style_names: List[str]) -> None:
    """
    Add a temporary paragraph for each built-in *paragraph* style so they're not 'latent'.
    DO NOT include table styles here if you're calling add_paragraph().
    """
    for style_name in style_names:
        print(f"[INFO] Ensuring style '{style_name}' exists")
        doc.add_paragraph("X", style=style_name)  # This raises an error if style_name is a table style
    # Returns None (modifies doc in-place)


def remove_temp_paragraphs(doc: DocxDocument) -> None:
    """
    Remove the dummy paragraphs that contain only 'X'.
    """
    to_remove = []
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "X":
            to_remove.append(i)
    for idx in reversed(to_remove):
        p = doc.paragraphs[idx]
        parent = p._element.getparent()
        parent.remove(p._element)


def override_built_in_style(
    doc: DocxDocument,
    base_name: str,
    custom_name: str,
    font_size: float,
    bold: bool,
    italic: bool,
    font_color: Tuple[int, int, int],
    space_before: float,
    space_after: float
) -> None:
    """
    Override a built-in paragraph style in doc.styles[base_name].
    Keep the original style_id to ensure Pandoc recognizes it.
    """
    try:
        style = doc.styles[base_name]  # type: BaseStyle
    except KeyError:
        print(f"[WARNING] style '{base_name}' not found. Skipping.")
        return

    if style.type != WD_STYLE_TYPE.PARAGRAPH:
        print(f"[WARNING] style '{base_name}' is not a paragraph style. Skipping.")
        return

    original_id = style.style_id

    # Rename visually but keep the underlying ID for Pandoc
    style.name = custom_name
    style.style_id = original_id

    # Font adjustments
    font = style.font
    font.name = "IBM Plex Sans"
    # Remove theme attributes
    rPr = font.element.rPr
    if rPr is not None and rPr.rFonts is not None:
        for theme_attrib in ("w:asciiTheme", "w:hAnsiTheme", "w:csTheme", "w:eastAsiaTheme"):
            rPr.rFonts.attrib.pop(qn(theme_attrib), None)

    font.size = Pt(font_size)
    font.bold = bold
    font.italic = italic

    r, g, b = font_color
    font.color.rgb = RGBColor(r, g, b)

    # Spacing adjustments
    pf = style.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)

def create_reference_doc(config_file: str, reference_docx: str) -> None:
    """
    1) Create a new .docx with python-docx
    2) Force built-in paragraph styles to appear
    3) Override each style from the YAML config
    4) Remove dummy paragraphs/tables
    5) Save as `reference_docx` (for Pandoc's --reference-doc)
    """
    doc = Document()

    # Only paragraph styles here
    needed_paragraph_styles = [
        "Normal", "Title", "Subtitle",
        "Heading 1", "Heading 2", "Heading 3",
        "Heading 4", "Heading 5", "Heading 6",
        "Quote"
    ]
    ensure_built_in_styles_exist(doc, needed_paragraph_styles)

    # Read config
    with open(config_file, "r", encoding="utf-8") as f:
        data = yaml.safe_load(f)

    # Apply overrides
    for style_def in data.get("styles", []):
        base_name = style_def["base_name"]
        custom_name = style_def["custom_name"]

        # Paragraph style
        font_size = style_def["font_size"]
        bold = style_def.get("bold", False)
        italic = style_def.get("italic", False)
        font_color = tuple(style_def.get("font_color", [0, 0, 0]))
        space_before = style_def.get("space_before", 0)
        space_after = style_def.get("space_after", 0)

        override_built_in_style(
            doc,
            base_name=base_name,
            custom_name=custom_name,
            font_size=font_size,
            bold=bold,
            italic=italic,
            font_color=font_color,
            space_before=space_before,
            space_after=space_after
        )

    # Cleanup
    remove_temp_paragraphs(doc)

    # Save
    doc.save(reference_docx)
    print(f"[INFO] Created reference doc: {reference_docx}")


def convert_md_to_word(input_md: str, output_docx: str, reference_docx: str) -> None:
    """
    Runs Pandoc with --reference-doc to generate the final .docx.
    Also uses 'markdown+footnotes+mark' to enable footnotes & ==highlight==.
    """
    cmd = [
        "pandoc",
        "--from=markdown+footnotes+mark",
        input_md,
        f"--reference-doc={reference_docx}",
        "-o",
        output_docx
    ]
    print(f"[INFO] Running: {' '.join(cmd)}")
    subprocess.run(cmd, check=True)
    print(f"[INFO] Converted {input_md} -> {output_docx} using {reference_docx}")


def main():
    """
    Usage:
       python format_doc.py <yaml_config> <input_md> <output_docx>
    """
    if len(sys.argv) < 4:
        print("Usage: format_doc.py <yaml_config> <input_md> <output_docx>")
        sys.exit(1)

    config_file = sys.argv[1]
    input_md = sys.argv[2]
    output_docx = sys.argv[3]

    reference_doc = "reference.docx"
    create_reference_doc(config_file, reference_doc)
    convert_md_to_word(input_md, output_docx, reference_doc)
    # Example usage
    apply_table_style(output_docx, "Light Shading", output_docx)

if __name__ == "__main__":
    main()
