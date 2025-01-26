#!/usr/bin/env python3

import subprocess
import sys

import yaml
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn  # For removing theme attributes
from docx.shared import Pt, RGBColor


def ensure_built_in_styles_exist(doc, style_names):
    """
    Add a temporary paragraph for each built-in style so that
    those styles are not 'latent'.
    """
    for style_name in style_names:
        print(f"[INFO] Ensuring style '{style_name}' exists")
        doc.add_paragraph("X", style=style_name)
    return doc


def remove_temp_paragraphs(doc):
    """Remove the dummy paragraphs that contain only 'X'."""
    to_remove = []
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "X":
            to_remove.append(i)
    for idx in reversed(to_remove):
        p = doc.paragraphs[idx]
        p._element.getparent().remove(p._element)


def override_built_in_style(doc, base_name, custom_name,
                            font_size, bold, italic,
                            font_color, space_before, space_after):
    """
    Override a built-in style in `doc.styles[base_name]`.
    Keep style_id to ensure Pandoc recognizes it (e.g. "Heading 1", "Quote").
    """
    try:
        style = doc.styles[base_name]
    except KeyError:
        print(f"[WARNING] style '{base_name}' not found. Skipping.")
        return

    if style.type != WD_STYLE_TYPE.PARAGRAPH:
        print(f"[WARNING] style '{base_name}' is not a paragraph style. Skipping.")
        return

    original_id = style.style_id

    # Rename visually but keep the underlying ID for Pandoc
    style.name = custom_name
    style.style_id = original_id

    font = style.font
    font.name = "IBM Plex Sans"

    # Remove theme attributes so Word won't revert to theme fonts
    rPr = font.element.rPr
    if rPr is not None and rPr.rFonts is not None:
        for theme_attrib in ["w:asciiTheme", "w:hAnsiTheme", "w:csTheme", "w:eastAsiaTheme"]:
            rPr.rFonts.attrib.pop(qn(theme_attrib), None)

    font.size = Pt(font_size)
    font.bold = bold
    font.italic = italic

    r, g, b = font_color
    font.color.rgb = RGBColor(r, g, b)

    pf = style.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    
def apply_table_styles(doc):
    """
    Apply custom styling to all tables in the Word document.
    """
    for table in doc.tables:
        # Set table width
        for col in table.columns:
            for cell in col.cells:
                cell.width = None  # You can adjust this to a fixed value if needed

        # Add borders
        for row in table.rows:
            for cell in row.cells:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                borders = OxmlElement('w:tcBorders')

                for border_name in ["top", "left", "bottom", "right"]:
                    border = OxmlElement(f"w:{border_name}")
                    border.set("w:val", "single")  # Border style
                    border.set("w:sz", "4")  # Border thickness
                    border.set("w:space", "0")
                    border.set("w:color", "000000")  # Black border
                    borders.append(border)

                tcPr.append(borders)

        # Optional: Add shading for alternate rows
        for i, row in enumerate(table.rows):
            if i % 2 == 1:  # Apply shading to odd rows
                for cell in row.cells:
                    cell_shading = OxmlElement("w:shd")
                    cell_shading.set("w:fill", "F2F2F2")  # Light gray shading
                    cell._element.get_or_add_tcPr().append(cell_shading)



def create_reference_doc(config_file, reference_docx):
    """
    1) Create a fresh doc
    2) Force built-in styles to appear
    3) Override each built-in style from the YAML data
    4) Remove dummy paragraphs
    5) Save as reference.docx
    """
    doc = Document()

    # Typical built-in styles we want to ensure are not latent
    needed_styles = [
        "Normal", "Title", "Subtitle",
        "Heading 1", "Heading 2", "Heading 3",
        "Heading 4", "Heading 5", "Heading 6",
        "Quote"
    ]
    ensure_built_in_styles_exist(doc, needed_styles)

    # Load config from YAML
    with open(config_file, "r", encoding="utf-8") as f:
        data = yaml.safe_load(f)

    # Apply overrides
    for style_def in data.get("styles", []):
        base_name = style_def["base_name"]
        custom_name = style_def["custom_name"]

        font_size = style_def["font_size"]
        bold = style_def.get("bold", False)
        italic = style_def.get("italic", False)
        font_color = tuple(style_def.get("font_color", [0, 0, 0]))
        space_before = style_def.get("space_before", 0)
        space_after = style_def.get("space_after", 0)

        override_built_in_style(
            doc,
            base_name=base_name,
            custom_name=custom_name,
            font_size=font_size,
            bold=bold,
            italic=italic,
            font_color=font_color,
            space_before=space_before,
            space_after=space_after
        )
        

    # Apply table styles
    apply_table_styles(doc)

    # Remove dummy paragraphs
    remove_temp_paragraphs(doc)

    # Save
    doc.save(reference_docx)
    print(f"[INFO] Created reference doc: {reference_docx}")


def convert_md_to_word(input_md, output_docx, reference_docx):
    """
    Runs Pandoc with --reference-doc=reference_docx to generate final DOCX.
    Also enables footnotes and highlight syntax with markdown+footnotes+mark.
    """
    cmd = [
        "pandoc",
        "--from=markdown+footnotes+mark",  # Enable footnotes, ==highlight==, etc.
        input_md,
        f"--reference-doc={reference_docx}",
        "-o",
        output_docx
    ]
    print(f"[INFO] Running: {' '.join(cmd)}")
    subprocess.run(cmd, check=True)
    print(f"[INFO] Converted {input_md} -> {output_docx} using {reference_docx}")


def main():
    """
    Usage:
       python format_doc.py <yaml_config> <input_md> <output_docx>
    """
    if len(sys.argv) < 4:
        print("Usage: format_doc.py <yaml_config> <input_md> <output_docx>")
        sys.exit(1)

    config_file = sys.argv[1]
    input_md = sys.argv[2]
    output_docx = sys.argv[3]

    reference_doc = "reference.docx"

    # Create the reference doc with custom IBM styles
    create_reference_doc(config_file, reference_doc)

    # Convert the Markdown to DOCX with Pandoc
    convert_md_to_word(input_md, output_docx, reference_doc)


if __name__ == "__main__":
    main()
