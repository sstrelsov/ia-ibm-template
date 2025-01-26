#!/usr/bin/env python3

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor


def create_custom_style(styles, custom_name, base_name, font_size, 
                        bold=False, italic=False, font_color=(0, 0, 0),
                        space_before=0, space_after=0):
    """
    Helper function to create a single custom paragraph style based on a built-in style.
    """
    # 1. Create the style (paragraph type)
    custom_style = styles.add_style(custom_name, WD_STYLE_TYPE.PARAGRAPH)
    
    # 2. Base it on the built-in style
    custom_style.base_style = styles[base_name]
    
    # 3. Mark as a Quick Style so it shows in Word's Styles gallery
    custom_style.quick_style = True

    # 4. Customize the font
    font = custom_style.font
    font.name = "IBM Plex Sans"
    font.size = Pt(font_size)
    font.bold = bold
    font.italic = italic
    
    # Convert (R,G,B) into an RGBColor object
    r, g, b = font_color
    font.color.rgb = RGBColor(r, g, b)

    # 5. Customize paragraph spacing
    paragraph_format = custom_style.paragraph_format
    paragraph_format.space_before = Pt(space_before)
    paragraph_format.space_after = Pt(space_after)

def create_custom_styles_with_quick_access(document):
    """
    Create multiple custom IBM styles (Title, Subtitle, Normal, Heading 1–6).
    All are flagged as Quick Styles so they'll appear in Word's Styles Pane.
    """
    styles = document.styles

    # Define each style you want to create in a dictionary-like structure:
    # Feel free to tweak sizes, spacing, and other attributes as you see fit.
    style_configs = [
        {
            "custom_name": "IBM Title",
            "base_name": "Title",
            "font_size": 28,
            "bold": True,
            "italic": False,
            "font_color": (0, 0, 0),
            "space_before": 0,
            "space_after": 12
        },
        {
            "custom_name": "IBM Subtitle",
            "base_name": "Subtitle",
            "font_size": 14,
            "bold": False,
            "italic": True,
            "font_color": (50, 50, 50),
            "space_before": 0,
            "space_after": 6
        },
        {
            "custom_name": "IBM Normal",
            "base_name": "Normal",
            "font_size": 11,
            "bold": False,
            "italic": False,
            "font_color": (0, 0, 0),
            "space_before": 0,
            "space_after": 0
        },
    ]

    # Add IBM versions of Heading 1–6
    # Adjust sizes/bold as needed
    heading_sizes = [24, 18, 14, 12, 10, 10]  # Example sizes, tweak as desired
    for i, size in enumerate(heading_sizes, start=1):
        style_configs.append({
            "custom_name": f"IBM Heading {i}",
            "base_name": f"Heading {i}",
            "font_size": size,
            "bold": True,
            "italic": False,
            "font_color": (0, 0, 0),
            "space_before": 0,
            "space_after": 8
        })

    # Create each style in the document
    for config in style_configs:
        create_custom_style(
            styles,
            custom_name=config["custom_name"],
            base_name=config["base_name"],
            font_size=config["font_size"],
            bold=config.get("bold", False),
            italic=config.get("italic", False),
            font_color=config.get("font_color", (0, 0, 0)),
            space_before=config.get("space_before", 0),
            space_after=config.get("space_after", 0)
        )

    print("All custom IBM styles created and added to Quick Styles!")

def main():
    # Create a new document
    doc = Document()

    # Create IBM custom styles
    create_custom_styles_with_quick_access(doc)

    # Demonstrate usage of the custom styles
    doc.add_paragraph("This is the IBM Title", style="IBM Title")
    doc.add_paragraph("This is the IBM Subtitle", style="IBM Subtitle")
    doc.add_paragraph("This is the IBM Normal (body text).", style="IBM Normal")
    
    for i in range(1, 7):
        doc.add_paragraph(f"This is IBM Heading {i}", style=f"IBM Heading {i}")

    # Save the document
    doc.save("custom_quick_styles_all.docx")
    print("Saved 'custom_quick_styles_all.docx' with all IBM Quick Styles!")

if __name__ == '__main__':
    main()
